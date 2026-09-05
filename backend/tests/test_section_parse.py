import tempfile
import unittest
from pathlib import Path

from app.services import file_parser
from app.services.ai_service import _file_user_payload
from app.services.section_intent import content_for_extract
from app.services.section_model import Section, SectionImage
from app.services.section_parse import parse_to_sections, vision_images_to_send

_PNG = bytes.fromhex(
    "89504e470d0a1a0a0000000d49484452000000010000000108060000001f15c489"
    "0000000a49444154789c63000100000500010d0a2db40000000049454e44ae426082"
)


class SectionParseExcelTest(unittest.TestCase):
    def test_atp_glo_sections_keep_result_sheet(self):
        src = Path("/tmp/user-xlsx/file.xlsx")
        if not src.exists():
            src = Path("/workspace/backend/data/uploads/ea35ba945782.xlsx")
        if not src.exists():
            self.skipTest("user ATP-GLO xlsx not cached")
        info = file_parser.save_upload("20210511 A375 CD73 ATP-GLO 汇总.xlsx", src.read_bytes())
        self.addCleanup(lambda: Path(info["path"]).unlink(missing_ok=True))
        name_meta = Path(info["path"]).parent / f".name-{info['file_id']}"
        self.addCleanup(lambda: name_meta.unlink(missing_ok=True))
        secs = parse_to_sections(info["file_id"])
        titles = [s.title for s in secs]
        self.assertTrue(any(t == "Result sheet" for t in titles))
        text, used, mode, _ = content_for_extract(secs)
        self.assertEqual(mode, "result")
        self.assertIn("HW100003", text)
        self.assertIn("Result sheet", text)
        self.assertGreater(len(text), 200)
        result = next(s for s in secs if s.title == "Result sheet")
        self.assertTrue(result.images)
        self.assertTrue(all(not im.data for im in result.images))
        self.assertTrue(any("无法解码" in (im.note or "") for im in result.images))
        self.assertIn("未能解码", text)
        self.assertFalse(any(im.mime.startswith("image/") and im.data for im in result.images))

    def test_xlsx_png_attaches_to_sheet(self):
        from openpyxl import Workbook
        from openpyxl.drawing.image import Image as XLImage

        tmp = Path(tempfile.mkdtemp())
        png_path = tmp / "dot.png"
        png_path.write_bytes(_PNG)
        wb = Workbook()
        ws = wb.active
        ws.title = "Result sheet"
        ws["A1"] = "HW100003"
        ws.add_image(XLImage(str(png_path)), "B2")
        xlsx_path = tmp / "t.xlsx"
        wb.save(xlsx_path)
        info = file_parser.save_upload("t.xlsx", xlsx_path.read_bytes())
        self.addCleanup(lambda: Path(info["path"]).unlink(missing_ok=True))
        secs = parse_to_sections(info["file_id"])
        result = next(s for s in secs if s.title == "Result sheet")
        self.assertTrue(any(im.data and im.mime.startswith("image/") for im in result.images))
        sent = vision_images_to_send([result], "result")
        self.assertTrue(sent)
        payload = _file_user_payload("正文", None, sent)
        self.assertIsInstance(payload, list)
        self.assertEqual(payload[0]["type"], "text")
        self.assertEqual(payload[1]["type"], "image_url")
        self.assertIn("image/png", payload[1]["image_url"]["url"])


class VisionSelectTest(unittest.TestCase):
    def test_process_images_not_sent_on_full(self):
        png = SectionImage("a.png", "image/png", data=_PNG)
        result = Section("Result", "sheet", "HW", [png], "result")
        process = Section("Raw", "sheet", "peak", [png], "process")
        self.assertEqual(len(vision_images_to_send([result], "result")), 1)
        self.assertEqual(vision_images_to_send([result, process], "full"), [])
        pic = Section("shot.png", "heading", "（整份为图片）", [png])
        self.assertEqual(len(vision_images_to_send([pic], "full")), 1)
        unread = SectionImage("a.emf", "application/octet-stream", note=".emf 暂无法解码")
        self.assertEqual(vision_images_to_send([Section("R", "sheet", "x", [unread])], "result"), [])


class DocxSectionTest(unittest.TestCase):
    def test_heading_keeps_nearby_image(self):
        from docx import Document
        from docx.shared import Inches

        tmp = Path(tempfile.mkdtemp())
        png_path = tmp / "dot.png"
        png_path.write_bytes(_PNG)
        doc = Document()
        doc.add_heading("方法", level=1)
        doc.add_paragraph("protocol")
        doc.add_heading("结果", level=1)
        doc.add_paragraph("IC50 12")
        doc.add_picture(str(png_path), width=Inches(0.2))
        docx_path = tmp / "t.docx"
        doc.save(docx_path)
        info = file_parser.save_upload("t.docx", docx_path.read_bytes())
        self.addCleanup(lambda: Path(info["path"]).unlink(missing_ok=True))
        text = file_parser.parse_to_text(info["file_id"])
        self.assertIn("结果", text)
        self.assertIn("IC50 12", text)
        secs = parse_to_sections(info["file_id"])
        titles = [s.title for s in secs]
        self.assertIn("结果", titles)
        result = next(s for s in secs if s.title == "结果")
        self.assertTrue(any(im.data for im in result.images))


if __name__ == "__main__":
    unittest.main()
